import io
from pathlib import Path
from reportlab.pdfgen import canvas
from docx import Document
from app.domain.extraction_models import ExtractionElementType
from app.services.extraction import ExtractionError, ExtractionService, PlainTextExtractor
from app.services.extraction_analysis import analyze_extraction


def make_pdf(path: Path):
    pdf = canvas.Canvas(str(path)); pdf.drawString(72, 720, "BEGINNING_SENTINEL"); pdf.drawString(72, 690, "MIDDLE_SENTINEL"); pdf.drawString(72, 660, "FINAL_PAGE_SENTINEL"); pdf.save()


def make_docx(path: Path):
    document = Document(); document.add_heading("DOCX_HEADING", level=1); document.add_paragraph("DOCX_BEGINNING_SENTINEL"); document.add_paragraph("DOCX_FINAL_SENTINEL"); document.save(path)
    table = document.add_table(rows=2, cols=2); table.cell(0, 0).text = "TABLE_CELL_A1"; table.cell(0, 1).text = "TABLE_CELL_B1"; table.cell(1, 0).text = "TABLE_CELL_A2"; table.cell(1, 1).text = "TABLE_FINAL_CELL"; document.save(path)


def test_txt_preserves_order_unicode_and_legitimate_repetition(tmp_path: Path):
    source = tmp_path / "source.txt"
    source.write_text("BEGINNING_SENTINEL\nREPEATED_SENTINEL\nMIDDLE_SENTINEL\nREPEATED_SENTINEL\nFINAL_SENTINEL\nUnicode café", encoding="utf-8")
    document = PlainTextExtractor().extract(source, "file-1", source.name, "text/plain")
    assert [element.text for element in document.elements] == ["BEGINNING_SENTINEL", "REPEATED_SENTINEL", "MIDDLE_SENTINEL", "REPEATED_SENTINEL", "FINAL_SENTINEL", "Unicode café"]
    assert document.statistics.word_count == 7
    assert document.statistics.element_count == 6
    assert document.validation.valid is True


def test_txt_empty_is_rejected(tmp_path: Path):
    source = tmp_path / "empty.txt"; source.write_text("", encoding="utf-8")
    try: PlainTextExtractor().extract(source, "file-1", source.name, "text/plain")
    except ExtractionError as exc: assert exc.code == "EMPTY_EXTRACTION"
    else: raise AssertionError("empty extraction should fail")


def test_pdf_docx_real_docling_extraction_preserves_sentinels(tmp_path: Path):
    pdf = tmp_path / "fixture.pdf"; docx = tmp_path / "fixture.docx"; make_pdf(pdf); make_docx(docx)
    service = ExtractionService()
    pdf_document = service.extract(pdf, "pdf-1", pdf.name, "application/pdf")
    docx_document = service.extract(docx, "docx-1", docx.name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    pdf_text = "\n".join(element.text for element in pdf_document.elements)
    docx_text = "\n".join(element.text for element in docx_document.elements)
    assert all(value in pdf_text for value in ("BEGINNING_SENTINEL", "MIDDLE_SENTINEL", "FINAL_PAGE_SENTINEL"))
    assert all(value in docx_text for value in ("DOCX_HEADING", "DOCX_BEGINNING_SENTINEL", "DOCX_FINAL_SENTINEL"))
    assert pdf_document.statistics.page_count == 1
    assert any(element.type == ExtractionElementType.HEADING for element in docx_document.elements)
    table_elements = [element for element in docx_document.elements if element.type == ExtractionElementType.TABLE]
    assert table_elements and "TABLE_FINAL_CELL" in table_elements[0].text


def test_analysis_is_deterministic(tmp_path: Path):
    source = tmp_path / "analysis.txt"; source.write_text("TITLE:\nparagraph text", encoding="utf-8")
    document = PlainTextExtractor().extract(source, "file-1", source.name, "text/plain")
    assert analyze_extraction(document) == analyze_extraction(document)
    assert analyze_extraction(document)["sectionCount"] == 1


def test_corrupt_pdf_maps_to_truthful_error(tmp_path: Path):
    source = tmp_path / "broken.pdf"; source.write_bytes(b"not a pdf")
    try: ExtractionService().extract(source, "file-1", source.name, "application/pdf")
    except ExtractionError as exc: assert exc.code == "CORRUPT_DOCUMENT"
    else: raise AssertionError("corrupt PDF should fail")
